"""Extraction: turn source files into plain text or structured JSON.

Dispatch by extension:

* ``.md`` / ``.txt``          — passthrough
* ``.pdf``                    — PyMuPDF text; pages without a text layer are
                                OCR'd (ocrmac / macOS Vision, tesseract fallback)
* ``.docx``                   — python-docx; ``.doc`` via macOS ``textutil``
* ``.xlsx`` / ``.xls`` / ``.csv`` — structured JSON (per-sheet headers + rows).
                                Los ``.xls`` se despachan por contenido, no por
                                extension: muchos exports de banca y
                                contabilidad son en realidad tablas HTML (o un
                                .xlsx) con el nombre cambiado.
* images (.png/.jpg/.heic...) — OCR
* code files                  — skipped ("code goes to codebase-memory-mcp")

Output goes to ``<extracted_dir>/<doc_id>.txt`` (or ``.json``).
"""

from __future__ import annotations

import csv
import io
import json
import logging
import shutil
import subprocess
import tempfile
from html.parser import HTMLParser
from pathlib import Path

log = logging.getLogger("brain")

TEXT_EXTS = {".md", ".txt", ".markdown", ".rst", ".org"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".heic", ".tiff", ".tif", ".bmp", ".gif", ".webp"}
CODE_EXTS = {
    ".py", ".js", ".ts", ".tsx", ".jsx", ".mjs", ".cjs", ".go", ".rs", ".java",
    ".c", ".cc", ".cpp", ".h", ".hpp", ".rb", ".sh", ".bash", ".zsh", ".swift",
    ".kt", ".kts", ".php", ".sql", ".scala", ".lua", ".pl", ".r", ".m", ".mm",
    ".toml", ".yaml", ".yml", ".ini", ".cfg", ".lock", ".ipynb",
}
CODE_SKIP_REASON = "code goes to codebase-memory-mcp"


class SkipFile(Exception):
    """File should be marked skipped (reason in str(exc))."""


class ExtractError(Exception):
    """Extraction failed."""


#: Archivos de bloqueo que Office deja abiertos junto al original ("~$x.xlsx").
#: No son documentos: son basura temporal de unos pocos bytes.
def es_temporal_de_office(path: Path) -> bool:
    return path.name.startswith("~$")


def esta_en_la_nube(path: Path) -> bool:
    """Contenido evictado a iCloud: figura con tamaño pero no ocupa bloques.

    Importa detectarlo aqui y no solo en scan: al abrirlo, PyMuPDF responde
    "Failed to open file as type pdf" y python-docx "Package not found", que
    parecen archivos corruptos y no lo son.
    """
    try:
        st = path.stat()
    except OSError:
        return False
    return st.st_size > 0 and st.st_blocks == 0


def extract_file(path: Path) -> tuple[str, str]:
    """Extract ``path``. Returns ``(content, kind)`` with kind ``text``/``json``.

    Raises :class:`SkipFile` or :class:`ExtractError`.
    """
    ext = path.suffix.lower()
    if es_temporal_de_office(path):
        raise SkipFile("archivo temporal de Office (~$), no es un documento")
    if esta_en_la_nube(path):
        raise SkipFile(
            "está en iCloud y no en el disco; descárgalo (brctl download) y vuelve a correr"
        )
    if ext in CODE_EXTS:
        raise SkipFile(CODE_SKIP_REASON)
    if ext in TEXT_EXTS:
        return path.read_text(encoding="utf-8", errors="replace"), "text"
    if ext == ".pdf":
        return _extract_pdf(path), "text"
    if ext == ".docx":
        return _extract_docx(path), "text"
    if ext == ".doc":
        return _extract_doc(path), "text"
    if ext in (".xlsx", ".xlsm"):
        return _extract_xlsx(path), "json"
    if ext == ".xls":
        return _extract_xls(path), "json"
    if ext == ".csv":
        return _extract_csv(path), "json"
    if ext in IMAGE_EXTS:
        return _ocr_image(path), "text"
    raise SkipFile(f"unsupported extension {ext or '(none)'}")


# -- OCR ---------------------------------------------------------------------


def ocr_available() -> bool:
    try:
        import ocrmac  # noqa: F401

        return True
    except ImportError:
        return shutil.which("tesseract") is not None


def _ocr_image(path: Path) -> str:
    try:
        from ocrmac import ocrmac as _ocrmac

        annotations = _ocrmac.OCR(str(path), recognition_level="accurate").recognize()
        return "\n".join(a[0] for a in annotations)
    except ImportError:
        pass
    if shutil.which("tesseract"):
        proc = subprocess.run(
            ["tesseract", str(path), "stdout", "-l", "spa+eng"],
            capture_output=True,
            text=True,
            timeout=300,
        )
        if proc.returncode == 0:
            return proc.stdout
        raise ExtractError(f"tesseract failed: {proc.stderr.strip()[:200]}")
    raise ExtractError("no OCR backend (install ocrmac or tesseract)")


# -- PDF ---------------------------------------------------------------------


def _extract_pdf(path: Path) -> str:
    import fitz  # PyMuPDF

    parts: list[str] = []
    with fitz.open(path) as doc:
        for pageno, page in enumerate(doc):
            text = page.get_text().strip()
            if not text:
                text = _ocr_pdf_page(page, pageno, path)
            parts.append(text)
    return "\n\n".join(parts).strip()


def _ocr_pdf_page(page, pageno: int, path: Path) -> str:
    if not ocr_available():
        log.warning("%s p.%d has no text layer and OCR unavailable", path.name, pageno + 1)
        return ""
    pix = page.get_pixmap(dpi=200)
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    try:
        pix.save(str(tmp_path))
        return _ocr_image(tmp_path)
    finally:
        tmp_path.unlink(missing_ok=True)


# -- Word --------------------------------------------------------------------


def _extract_docx(path: Path) -> str:
    import docx

    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n\n".join(p for p in parts if p.strip())


def _extract_doc(path: Path) -> str:
    if not shutil.which("textutil"):
        raise ExtractError(".doc extraction requires macOS textutil")
    with tempfile.TemporaryDirectory() as td:
        out = Path(td) / (path.stem + ".txt")
        proc = subprocess.run(
            ["textutil", "-convert", "txt", "-output", str(out), str(path)],
            capture_output=True,
            text=True,
            timeout=120,
        )
        if proc.returncode != 0 or not out.exists():
            raise ExtractError(f"textutil failed: {proc.stderr.strip()[:200]}")
        return out.read_text(encoding="utf-8", errors="replace")


# -- Tabular -----------------------------------------------------------------


def _extract_xlsx(path: Path) -> str:
    import openpyxl

    # Se abre como stream y no por ruta a proposito: load_workbook(str(path))
    # rechaza el archivo por la EXTENSION antes de mirarlo, asi que un .xlsx
    # renombrado a .xls (caso comun en exports) fallaria pese a ser valido.
    with path.open("rb") as fh:
        wb = openpyxl.load_workbook(fh, read_only=True, data_only=True)
        sheets = _hojas_openpyxl(wb)
    return _sheets_json(path, sheets)


def _hojas_openpyxl(wb) -> list[dict]:
    sheets = []
    for ws in wb.worksheets:
        rows = [
            ["" if c is None else str(c) for c in row]
            for row in ws.iter_rows(values_only=True)
        ]
        headers = rows[0] if rows else []
        sheets.append({"sheet": ws.title, "headers": headers, "rows": rows[1:]})
    wb.close()
    return sheets


def _sheets_json(path: Path, sheets: list[dict]) -> str:
    return json.dumps({"source": path.name, "sheets": sheets}, ensure_ascii=False, indent=1)


def _extract_xls(path: Path) -> str:
    """Extrae un ``.xls`` despachando por los primeros bytes, no por el nombre.

    La extension .xls miente a menudo: los exports de cartolas bancarias y de
    software contable suelen ser una tabla HTML (o directamente un .xlsx) con el
    nombre cambiado, y abrirlos con xlrd falla con "Unsupported format". Se mira
    la firma real del archivo:

    * ``D0 CF 11 E0`` — OLE2, el .xls binario de verdad  -> xlrd
    * ``PK``          — zip, o sea un .xlsx mal nombrado -> openpyxl
    * cualquier otra  — se intenta como HTML con <table>
    """
    firma = path.read_bytes()[:8]
    if firma.startswith(b"\xd0\xcf\x11\xe0"):
        return _extract_xls_ole2(path)
    if firma.startswith(b"PK"):
        return _extract_xlsx(path)
    return _extract_html_tables(path)


def _extract_xls_ole2(path: Path) -> str:
    """.xls binario (Excel 97-2003) via xlrd."""
    import xlrd

    wb = xlrd.open_workbook(str(path))
    sheets = []
    for ws in wb.sheets():
        rows = []
        for r in range(ws.nrows):
            fila = []
            for c in range(ws.ncols):
                celda = ws.cell(r, c)
                # Las fechas viajan como float; sin convertirlas el grafo
                # recibiria "45231.0" en vez de una fecha (regla de oro #1).
                if celda.ctype == xlrd.XL_CELL_DATE:
                    try:
                        fila.append(xlrd.xldate_as_datetime(celda.value, wb.datemode).isoformat())
                        continue
                    except (ValueError, OverflowError):
                        pass
                fila.append("" if celda.value is None else str(celda.value))
            rows.append(fila)
        sheets.append({"sheet": ws.name, "headers": rows[0] if rows else [], "rows": rows[1:]})
    return _sheets_json(path, sheets)


class _TablasHTML(HTMLParser):
    """Saca las filas de cada <table> de un documento HTML."""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.tablas: list[list[list[str]]] = []
        self._tabla: list[list[str]] | None = None
        self._fila: list[str] | None = None
        self._celda: list[str] | None = None

    def handle_starttag(self, tag, attrs):
        if tag == "table":
            self._tabla = []
        elif tag == "tr" and self._tabla is not None:
            self._fila = []
        elif tag in ("td", "th") and self._fila is not None:
            self._celda = []

    def handle_endtag(self, tag):
        if tag in ("td", "th") and self._celda is not None:
            self._fila.append(" ".join("".join(self._celda).split()))
            self._celda = None
        elif tag == "tr" and self._fila is not None:
            if any(v for v in self._fila):
                self._tabla.append(self._fila)
            self._fila = None
        elif tag == "table" and self._tabla is not None:
            if self._tabla:
                self.tablas.append(self._tabla)
            self._tabla = None

    def handle_data(self, data):
        if self._celda is not None:
            self._celda.append(data)


def _extract_html_tables(path: Path) -> str:
    """Tabla HTML disfrazada de hoja de calculo (cartolas, reportes contables)."""
    raw = path.read_bytes()
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            texto = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        texto = raw.decode("utf-8", errors="replace")

    parser = _TablasHTML()
    parser.feed(texto)
    if not parser.tablas:
        raise ExtractError(
            "el .xls no es OLE2 ni zip ni trae tablas HTML; formato desconocido"
        )
    sheets = [
        {
            "sheet": f"tabla_{i + 1}" if len(parser.tablas) > 1 else path.stem,
            "headers": t[0],
            "rows": t[1:],
        }
        for i, t in enumerate(parser.tablas)
    ]
    return _sheets_json(path, sheets)


def _extract_csv(path: Path) -> str:
    raw = path.read_text(encoding="utf-8-sig", errors="replace")
    try:
        dialect = csv.Sniffer().sniff(raw[:4096], delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel
    rows = [row for row in csv.reader(io.StringIO(raw), dialect)]
    headers = rows[0] if rows else []
    payload = {
        "source": path.name,
        "sheets": [{"sheet": path.stem, "headers": headers, "rows": rows[1:]}],
    }
    return json.dumps(payload, ensure_ascii=False, indent=1)
